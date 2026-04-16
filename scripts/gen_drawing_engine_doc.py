"""Generate DOCX document explaining the Technical Drawing Engine."""

from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

OUTPUT = Path(__file__).parent.parent / "output" / "Technical_Drawing_Engine.docx"
OUTPUT.parent.mkdir(parents=True, exist_ok=True)

doc = Document()

# -- Styles --
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)
style.paragraph_format.space_after = Pt(6)

for level in range(1, 4):
    hs = doc.styles[f"Heading {level}"]
    hs.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)

# -- Title --
title = doc.add_heading("Technical Drawing Engine — Escora.AI", level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph(
    "Architecture overview, module breakdown, and improvement roadmap "
    "for the ABNT/NBR-compliant DXF generation engine."
).alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph("")

# ===================================================================
doc.add_heading("How the Engine Works", level=1)

doc.add_paragraph(
    "The engine has 4 layers, bottom to top. Each layer builds on the "
    "previous one, separating concerns between domain knowledge (NBR standards), "
    "low-level DXF writing, high-level drawing orchestration, and view generation."
)

# -- Layer 1 --
doc.add_heading("1. nbr.py — Standards Database", level=2)

doc.add_paragraph(
    "A static registry of all ABNT/NBR rules encoded as Python enums and dataclasses. "
    "It knows nothing about DXF — it is pure domain knowledge:"
)

items = [
    ("SheetFormat", "A0–A4 sizes, margins, legend widths (NBR 10068)"),
    ("LineType", "8 line types A through K with dash patterns and wide/narrow classification (NBR 8403)"),
    ("Scale", "Bidirectional real↔paper conversion, auto text height calculation (NBR 8196)"),
    ("DimensionRules", "Gaps, overshoots, terminator types, method 1 vs 2 (NBR 10126)"),
    ("HatchMaterial", "Pattern names per material: metal, concrete, brick, wood, etc. (NBR 12298)"),
    ("ViewArrangement", "1st/3rd diedro view positioning grid (NBR 10067)"),
]
for name, desc in items:
    p = doc.add_paragraph(style="List Bullet")
    run_b = p.add_run(name)
    run_b.bold = True
    p.add_run(f" — {desc}")

doc.add_paragraph(
    "Everything downstream references this module instead of hardcoding magic numbers."
)

# -- Layer 2 --
doc.add_heading("2. primitives.py — Low-Level DXF Generators", level=2)

doc.add_paragraph(
    "Functions that receive an ezdxf modelspace and draw individual elements:"
)

items2 = [
    ("Dimensions", "add_linear_dimension, add_chain_dimensions, add_radius_dimension, "
     "add_diameter_dimension. Each creates an ezdxf dimension entity with the NBR_DIM style "
     "(arrow size, text position, decimal comma, etc.)"),
    ("Hatching", "add_hatch fills a polygon with a material pattern. add_section_hatch "
     "auto-rotates angles for adjacent parts (NBR 12298 rule)."),
    ("Text", "add_text with 9-point alignment, add_room_label (name + area), "
     "add_leader with arrowhead."),
    ("Section indicators", "add_cutting_plane draws the dash-dot line with direction "
     "arrows and labels."),
]
for name, desc in items2:
    p = doc.add_paragraph(style="List Bullet")
    run_b = p.add_run(name)
    run_b.bold = True
    p.add_run(f": {desc}")

doc.add_paragraph(
    "These are stateless — they don't know about sheets, scales, or layers. "
    "They just write ezdxf entities."
)

# -- Layer 3 --
doc.add_heading("3. sheet.py — TechnicalSheet (Main Entry Point)", level=2)

doc.add_paragraph(
    "The facade class that ties everything together. On creation it:"
)

steps = [
    "Creates an ezdxf.new(\"R2010\") document.",
    "Registers custom linetypes (DASHED, DASHDOT, DASHDOTDOT) — ezdxf doesn't include them by default.",
    "Creates ~14 layers with correct colors, lineweights, and linetypes (ARCH or STRUCT set).",
    "Registers the NBR_DIM dimension style with scale-aware text heights.",
    "Registers the \"NBR\" text style using isocpeur.shx (ISO technical font).",
]
for i, step in enumerate(steps, 1):
    doc.add_paragraph(f"{i}. {step}", style="List Number")

doc.add_paragraph(
    "Then it exposes high-level methods: draw_wall, draw_door, draw_window, "
    "add_dimension, add_title_block, etc. — each delegating to primitives with "
    "the correct layer, line type, and scale context."
)

# -- Layer 4 --
doc.add_heading("4. views.py + perspectives.py — View Generators", level=2)

doc.add_paragraph(
    "views.py: Solid3D → orthographic multi-view projections. Each 3D face is projected "
    "to the chosen view plane using axis selection + normal dot-product visibility check. "
    "Faces project as visible (Type A continuous) or hidden (Type E dashed). "
    "generate_section_from_walls cuts walls with a plane and draws hatched cut faces "
    "vs visible background faces."
)
doc.add_paragraph(
    "perspectives.py: iso_project(x,y,z) maps to 2D isometric (30° axes), cav_project "
    "maps to cavaleira (front true + depth reduced). draw_elevation projects walls/openings "
    "to a facade view."
)

# -- Data Flow --
doc.add_heading("Data Flow", level=2)

flow_lines = [
    "User input (walls, openings, rooms)",
    "    → TechnicalSheet (creates DXF doc, sets up layers/styles)",
    "        → draw_wall / draw_door / add_dimension (calls primitives)",
    "        → draw_orthographic_views / draw_isometric / generate_section",
    "            → primitives write ezdxf entities",
    "    → .save(\"output.dxf\")",
]
for line in flow_lines:
    p = doc.add_paragraph(line)
    p.paragraph_format.space_after = Pt(0)
    for run in p.runs:
        run.font.name = "Consolas"
        run.font.size = Pt(10)

doc.add_paragraph("")

# ===================================================================
doc.add_heading("What's Weak / How to Improve", level=1)

improvements = [
    (
        "1. No Real 3D Model — Everything is Disconnected",
        "Right now, walls are just (p1, p2, thickness) tuples passed independently to each function. "
        "The floor plan, section, elevation, and isometric are generated from separate calls with "
        "duplicated data. There is no single source of truth.",
        "A proper BuildingModel class that stores walls, openings, slabs, roof as a unified 3D model. "
        "Then each view generator reads from it: sheet.draw_plan(model), sheet.draw_section(model, cut_plane), etc.",
    ),
    (
        "2. Wall Intersections Not Resolved",
        "draw_wall draws each wall as an independent rectangle. Where walls meet at corners, the "
        "rectangles overlap — a real drawing needs proper T-junctions and L-junctions where wall "
        "outlines merge into clean corners.",
        "This is the biggest visual quality gap vs professional output. Requires computational geometry "
        "to detect intersections and merge polylines.",
    ),
    (
        "3. Doors/Windows Don't Cut the Wall",
        "draw_door draws the arc symbol but doesn't actually create a gap in the wall fill. In a real "
        "planta baixa, the wall hatch/fill must be interrupted where openings exist.",
        "Requires wall-aware opening placement — the wall geometry must be split around each opening.",
    ),
    (
        "4. No Roof Drawing",
        "Sections and elevations show flat ceiling slabs but no actual roof geometry (telhado). A real "
        "corte shows the roof structure (tesouras, caibros, ripas, telhas). The fachada shows the roof "
        "profile and overhang (beiral).",
        "Add roof primitives: gable, hip, shed profiles with configurable slope and overhang.",
    ),
    (
        "5. Section Generator is Simplified",
        "generate_section_from_walls only handles straight horizontal cutting planes and doesn't support "
        "corte em desvio (offset sections), meio-corte (half sections), door/window openings in section, "
        "floor level changes (desníveis), or foundation detail below the floor slab.",
        "Implement offset cutting planes and opening-aware section generation.",
    ),
    (
        "6. Isometric Only Handles Axis-Aligned Boxes",
        "draw_isometric_from_walls snaps all walls to cardinal directions. Diagonal walls, curved walls, "
        "or non-rectangular elements can't be drawn. The isometric circle approximation uses a polyline "
        "instead of the proper 4-arc construction (which looks better in AutoCAD).",
        "Support arbitrary wall orientations via proper 3D→2D projection and implement the 4-arc ellipse method.",
    ),
    (
        "7. No Automatic Dimensioning",
        "Dimensions are manually placed. A professional tool should auto-detect wall endpoints, opening "
        "positions, and room dimensions, then generate chain dimensions with proper spacing, no duplicates, "
        "and correct NBR ordering (smaller dimensions closer to the object).",
        "Build an auto-dimensioner that analyzes the BuildingModel and places dimensions automatically.",
    ),
    (
        "8. No Staircase, Furniture, or Site Elements",
        "Missing common architectural symbols: stairs (with arrow + step count), furniture blocks (sink, "
        "toilet, shower, bed), north arrow, site plan elements (sidewalk, driveway, vegetation).",
        "Create a symbol library with parametric architectural blocks.",
    ),
    (
        "9. No DXF/Image Input (Sketch Reader)",
        "The original vision was \"read hand-drawn sketches → DXF.\" That requires image processing "
        "(OpenCV or similar) to detect lines, then classify them as walls/dimensions/annotations. "
        "This is the hardest piece and isn't started yet.",
        "Integrate OpenCV line detection + ML classification for sketch-to-model conversion.",
    ),
    (
        "10. No Multi-Story Support",
        "Everything assumes a single floor. A 2-story building needs coordinated plans per floor, "
        "multi-level sections, and stair connections.",
        "Extend BuildingModel with a floor/level hierarchy.",
    ),
    (
        "11. Title Block is Basic",
        "The current carimbo is a 5-row rectangle. Professional title blocks include: company logo area, "
        "revision table, reference drawing list, approval signatures, and follow specific company templates.",
        "Support custom title block templates loaded from DXF block definitions.",
    ),
    (
        "12. No Paper Space / Viewports",
        "Everything is drawn in model space. Professional DXFs use paper space (Layout tabs) with viewports "
        "at different scales — e.g., the plan at 1:50, a detail at 1:20, and the title block at 1:1, all "
        "on the same sheet.",
        "Implement ezdxf Layout/Viewport support for multi-scale sheets.",
    ),
]

for title_text, problem, fix in improvements:
    doc.add_heading(title_text, level=2)
    p_prob = doc.add_paragraph()
    run_label = p_prob.add_run("Problem: ")
    run_label.bold = True
    p_prob.add_run(problem)

    p_fix = doc.add_paragraph()
    run_label2 = p_fix.add_run("Fix: ")
    run_label2.bold = True
    run_label2.font.color.rgb = RGBColor(0x00, 0x6B, 0x3F)
    p_fix.add_run(fix)

# ===================================================================
doc.add_heading("Priority Ranking", level=1)

doc.add_paragraph(
    "If development continues, this is the recommended order based on impact:"
)

table = doc.add_table(rows=8, cols=3)
table.style = "Light Grid Accent 1"
table.alignment = WD_TABLE_ALIGNMENT.CENTER

headers = ["#", "Improvement", "Impact"]
for i, h in enumerate(headers):
    cell = table.rows[0].cells[i]
    cell.text = h
    for run in cell.paragraphs[0].runs:
        run.bold = True

priorities = [
    ("1", "Unified BuildingModel", "Enables everything else"),
    ("2", "Wall intersection resolution", "Visual quality jump"),
    ("3", "Openings cut through walls", "Essential for planta baixa"),
    ("4", "Auto-dimensioning", "Huge time saver"),
    ("5", "Roof geometry", "Completes sections + elevations"),
    ("6", "Paper space viewports", "Professional DXF output"),
    ("7", "Sketch/image input", "The differentiator feature"),
]

for row_idx, (num, imp, impact) in enumerate(priorities, 1):
    table.rows[row_idx].cells[0].text = num
    table.rows[row_idx].cells[1].text = imp
    table.rows[row_idx].cells[2].text = impact

# Set column widths
for row in table.rows:
    row.cells[0].width = Inches(0.4)
    row.cells[1].width = Inches(3.0)
    row.cells[2].width = Inches(2.5)

doc.add_paragraph("")

# ===================================================================
doc.add_heading("File Structure", level=1)

files = [
    ("src/drawing/__init__.py", "Public API, exports TechnicalSheet and NBR"),
    ("src/drawing/nbr.py", "All NBR constants — formats, lines, scales, text, dimensions, hatching, projections"),
    ("src/drawing/primitives.py", "Low-level DXF generators — dimensions, hatching, text, leaders, cutting planes"),
    ("src/drawing/sheet.py", "TechnicalSheet facade — layers, title block, walls, doors, windows, save"),
    ("src/drawing/views.py", "Orthographic projections, Solid3D model, section generation"),
    ("src/drawing/perspectives.py", "Isometric, cavaleira, elevation generators"),
    ("api/routes/drawing.py", "REST API: /floor-plan, /section, /perspective, /formats"),
    ("scripts/test_drawing_engine.py", "Test script generating 5 sample DXFs"),
]

file_table = doc.add_table(rows=len(files) + 1, cols=2)
file_table.style = "Light Grid Accent 1"
file_table.alignment = WD_TABLE_ALIGNMENT.CENTER

file_table.rows[0].cells[0].text = "File"
file_table.rows[0].cells[1].text = "Purpose"
for run in file_table.rows[0].cells[0].paragraphs[0].runs:
    run.bold = True
for run in file_table.rows[0].cells[1].paragraphs[0].runs:
    run.bold = True

for i, (fname, purpose) in enumerate(files, 1):
    cell0 = file_table.rows[i].cells[0]
    cell0.text = fname
    for run in cell0.paragraphs[0].runs:
        run.font.name = "Consolas"
        run.font.size = Pt(9)
    file_table.rows[i].cells[1].text = purpose

for row in file_table.rows:
    row.cells[0].width = Inches(2.8)
    row.cells[1].width = Inches(3.5)

# -- Footer --
doc.add_paragraph("")
p_footer = doc.add_paragraph("Generated by Escora.AI — Technical Drawing Engine")
p_footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_footer.runs[0].font.size = Pt(9)
p_footer.runs[0].font.color.rgb = RGBColor(0x88, 0x88, 0x88)

# -- Save --
doc.save(str(OUTPUT))
print(f"Saved: {OUTPUT}")
